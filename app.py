from flask import Flask, render_template, request, jsonify
from main import retrieval_qa_pipline, handle_greetings
app = Flask(__name__, static_folder='static',template_folder='templates')
import os
import logging
from docx import Document
import re
from flask_mail import Mail, Message
from flask import Flask
import io


# Configure Flask Mail for Outlook.com
app.config['MAIL_SERVER'] = 'smtp-mail.outlook.com'
app.config['MAIL_PORT'] = 587  # Port 587 for STARTTLS
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = 'sasipriya004@outlook.com'  # Enter your Outlook.com email address
app.config['MAIL_PASSWORD'] = 'Karu14399!'  # Enter your Outlook.com email password
app.config['MAIL_DEFAULT_SENDER'] = 'sasipriya004@outlook.com'  # Enter your Outlook.com email address

mail = Mail(app)


from constants import (
   
    MODELS_PATH,

)

qa = retrieval_qa_pipline("cpu")
chat_history = []
def main(query):
    global chat_history

    if query == "exit":
        answer = "Goodbye!"
    elif handle_greetings(query):
        answer = "Hello! I am AthenaHR, here to assist you with any queries you may have related to Human resources within the company."
    else:
        # If chat history is not empty, use it for processing
        # if chat_history:
        #     res = qa({"question": query, "chat_history": chat_history})
        # else:
        # If chat history is empty, initialize it and make the first query
        res = qa(query)

        # Extract the result
        answer, docs = res["result"], res["source_documents"]

        if answer is not None:
            formatted_answer = format_answer(answer)
            chat_history.append({"question": query, "answer": formatted_answer})
        else:
            # Handle the case when 'result' key is not present in the response
            answer = "No answer found"

    return answer, chat_history

def format_answer(answer):
    lines = answer.split('\n')
    list_items = []
    additional_content_before = []
    additional_content_after = []

    in_list_content = False
    for line in lines:
        match = re.match(r'(\d+\.\s)(.+)', line)
        if match:
            in_list_content = True
            number = match.group(1)
            description = match.group(2).strip()
            list_items.append(f'<li>{number}{description}</li>')
        elif in_list_content:
            additional_content_after.append(line.strip())
        else:
            additional_content_before.append(line.strip())

    if list_items:
        list_content = f'<ul>{"".join(list_items)}</ul>'
    else:
        list_content = ''

    # Combine all content in the correct order
    additional_content_before = " ".join(additional_content_before)
    additional_content_after = " ".join(additional_content_after)

    formatted_output = f'<p>Analyst: {additional_content_before}</p>\n{list_content}'
    
    if additional_content_after:
        formatted_output += f'\n<p>{additional_content_after}</p>'

    return formatted_output

@app.route('/submit', methods=['POST'])
def handle_submit():
    query = request.form.get('user_input')
    answer, chat_history = main(query)
    print("chat:", chat_history)
    return render_template('Resumeanalyser.html', query=query, answer=answer, query_history=chat_history)

def save_chat_history_to_docx(email_hr_body):
    #with app.app_context():
        doc = Document()

        for entry in chat_history:
            doc.add_paragraph(f"User: {entry['question']}")
            doc.add_paragraph(f"Chatbot: {entry['answer']}")
            doc.add_paragraph("")  # Add an empty paragraph for spacing

        doc.save('chat_history.docx')
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0) 
        
        subject = 'Resume update'
        recipient = 'sasipriya.a@quadwave.com'  # Replace with the recipient's email address

        # Build the email body
        email_body = f"Good evening HR \n\n {email_hr_body}"

        msg = Message(subject=subject, recipients=[recipient], body=email_body)
        msg.attach("chat_history.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", doc_io.read())

        mail.send(msg)
    

@app.route('/save_history', methods=['POST'])
def save_history():
    email_hr_body = request.form.get('email_hr_body')
    save_chat_history_to_docx(email_hr_body)
    return jsonify(result="Email sent to HR successfully")

# Flask route for rendering the initial form
@app.route('/index')
def index():
    return render_template('index.html',)
@app.route('/resume')
def resume():
    return render_template('Resumeanalyser.html')


if __name__ == '__main__':
    app.run(debug=True, port=8000)